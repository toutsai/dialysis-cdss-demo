from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "progress_report_for_colleagues.md"
OUTPUT = ROOT / "docs" / "透析中心查房與藥物調整輔助系統_進度報告.docx"


ACCENT = RGBColor(30, 83, 159)
LIGHT_BLUE = "EAF2FF"
LIGHT_GRAY = "F4F6F8"
GRID = "D9E1EA"
TEXT = RGBColor(25, 35, 50)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = GRID) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    font = run.font
    font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    if size is not None:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if color is not None:
        font.color.rgb = color


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(5)

    for name, size, before, after in [
        ("Heading 1", 15, 12, 6),
        ("Heading 2", 12.5, 8, 4),
        ("Heading 3", 11, 6, 3),
    ]:
        style = styles[name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title_block(doc: Document, title: str, date_line: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    set_run_font(run, size=20, bold=True, color=ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("第一版雛形與後續工作整理")
    set_run_font(run, size=12, bold=False, color=RGBColor(90, 100, 115))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(14)
    run = meta.add_run(f"報告日期：{date_line.replace('日期：', '')}")
    set_run_font(run, size=10.5, bold=True, color=RGBColor(70, 85, 105))

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(10)
    run = note.add_run("用途：供透析中心同仁、護理長與資訊室初步了解目前進度、架構方向與後續試行準備。")
    set_run_font(run, size=10.5, color=RGBColor(65, 75, 90))


def add_code_block(doc: Document, lines: list[str]) -> None:
    for index, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_before = Pt(1 if index else 6)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(45, 55, 70)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=10.3, color=TEXT)


def build_docx() -> Path:
    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style_document(doc)

    title = source_lines[0].lstrip("# ").strip()
    date_line = source_lines[2].strip()
    add_title_block(doc, title, date_line)

    in_code = False
    code_lines: list[str] = []
    skip_until = 3
    for raw in source_lines[skip_until:]:
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, code_lines)
                in_code = False
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("- "):
            add_bullet(doc, stripped[2:].strip())
        else:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            set_run_font(run, size=10.5, color=TEXT)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("透析中心查房與藥物調整輔助系統｜內部討論草案")
    set_run_font(footer_run, size=8.5, color=RGBColor(120, 130, 145))

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_docx())
